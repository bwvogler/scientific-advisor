import io
import uuid
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from abc import ABC, abstractmethod

import PyPDF2
from docx import Document as DocxDocument
import aiofiles

from ...api.schemas.models import Document, DocumentType
from ...config.settings import settings

logger = logging.getLogger(__name__)

class DocumentProcessor(ABC):
    """Abstract base class for document processors."""
    
    @abstractmethod
    def can_process(self, filename: str) -> bool:
        """Check if this processor can handle the given file."""
        pass
    
    @abstractmethod
    def extract_text(self, file_content: bytes) -> str:
        """Extract text content from the file."""
        pass

class PDFProcessor(DocumentProcessor):
    """Process PDF documents."""
    
    def can_process(self, filename: str) -> bool:
        return filename.lower().endswith('.pdf')
    
    def extract_text(self, file_content: bytes) -> str:
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise ValueError(f"Failed to process PDF: {e}")

class DocxProcessor(DocumentProcessor):
    """Process DOCX documents."""
    
    def can_process(self, filename: str) -> bool:
        return filename.lower().endswith('.docx')
    
    def extract_text(self, file_content: bytes) -> str:
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError(f"Failed to process DOCX: {e}")

class TxtProcessor(DocumentProcessor):
    """Process plain text documents."""
    
    def can_process(self, filename: str) -> bool:
        return filename.lower().endswith('.txt')
    
    def extract_text(self, file_content: bytes) -> str:
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding).strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return file_content.decode('utf-8', errors='replace').strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from TXT: {e}")
            raise ValueError(f"Failed to process TXT: {e}")

class MarkdownProcessor(DocumentProcessor):
    """Process Markdown documents."""
    
    def can_process(self, filename: str) -> bool:
        return filename.lower().endswith(('.md', '.markdown'))
    
    def extract_text(self, file_content: bytes) -> str:
        try:
            return file_content.decode('utf-8').strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from Markdown: {e}")
            raise ValueError(f"Failed to process Markdown: {e}")

class DocumentIngestionService:
    """Service for processing and ingesting documents."""
    
    def __init__(self):
        self.processors = [
            PDFProcessor(),
            DocxProcessor(),
            TxtProcessor(),
            MarkdownProcessor()
        ]
    
    def process_document(self, filename: str, file_content: bytes, 
                        customer: Optional[str] = None, 
                        project: Optional[str] = None,
                        metadata: Optional[Dict[str, Any]] = None) -> Document:
        """Process a document and return a Document object."""
        try:
            # Find appropriate processor
            processor = None
            for proc in self.processors:
                if proc.can_process(filename):
                    processor = proc
                    break
            
            if not processor:
                raise ValueError(f"No processor available for file type: {filename}")
            
            # Extract text content
            content = processor.extract_text(file_content)
            
            if not content.strip():
                raise ValueError("Document appears to be empty or unreadable")
            
            # Determine document type
            doc_type = self._get_document_type(filename)
            
            # Create Document object
            document = Document(
                id=str(uuid.uuid4()),
                filename=filename,
                content=content,
                document_type=doc_type,
                customer=customer,
                project=project,
                date=datetime.utcnow(),
                metadata=metadata or {},
                file_size=len(file_content)
            )
            
            logger.info(f"Processed document: {filename} ({len(content)} characters)")
            return document
            
        except Exception as e:
            logger.error(f"Failed to process document {filename}: {e}")
            raise
    
    async def save_document_file(self, document: Document, file_content: bytes) -> str:
        """Save the original file to disk."""
        try:
            file_path = f"{settings.chroma_db_path}/../documents/{document.id}_{document.filename}"
            
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            logger.info(f"Saved document file: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Failed to save document file: {e}")
            raise
    
    def _get_document_type(self, filename: str) -> DocumentType:
        """Determine document type from filename."""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return DocumentType.PDF
        elif filename_lower.endswith('.docx'):
            return DocumentType.DOCX
        elif filename_lower.endswith('.txt'):
            return DocumentType.TXT
        elif filename_lower.endswith(('.md', '.markdown')):
            return DocumentType.MD
        else:
            # Default to TXT for unknown types
            return DocumentType.TXT
